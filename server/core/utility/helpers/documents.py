import asyncio
import base64
import re
from difflib import Differ
from io import BytesIO
from pathlib import Path
from typing import List

import httpx
from docusign_esign import EnvelopesApi, EnvelopeDefinition, TemplateRole, ApiClient, Document, Signer, CarbonCopy, \
    Recipients
from fastapi import HTTPException
from langchain_core.messages import SystemMessage, HumanMessage
from reportlab.pdfgen import canvas
from tavily import AsyncTavilyClient

from config import BaseConfig
from core.llm.openai import llm
from core.utility.constants import APIScope, SignStatus
from core.utility.prompts.documents import magic_edit_instructions, insight_extraction_instructions, \
    report_planner_query_writer_instructions, report_planner_instructions, insight_writer_instructions, \
    query_writer_instructions
from schemas.documents import SectionState, AIEdit, ReportState, Extractions, Queries, Insights, InsightState, SignEmail
from schemas.users import User

settings = BaseConfig()

async def get_access_code(code):
    # Exchange the authorization code for an access token
    TOKEN_URL = "https://account-d.docusign.com/oauth/token"
    auth_key = f"{settings.INTEGRATION_KEY}:{settings.CLIENT_SECRET}"
    encoded_auth = base64.b64encode(auth_key.encode("ascii")).decode("ascii")

    async with httpx.AsyncClient() as client:
        response = await client.post(
            TOKEN_URL,
            headers={
                "Authorization": f"Basic {encoded_auth}",
                "Accept": "application/json"
            },
            data={
                "grant_type": "authorization_code",
                "code": code,
            },
        )
        token_data = response.json()

        return token_data

    return {}

def build_redirect_url(is_esign):
    scope = APIScope.ESIGNATURE
    redirect_url = settings.SIGN_REDIRECT_URL
    if not is_esign:
        scope = APIScope.NAVIGATOR.value
        redirect_url = settings.NAV_REDIRECT_URL

    # Redirect user to third-party authorization endpoint
    # TODO: set the base url as an env var (N.B: This is the dev-endpoint)
    AUTHORIZE_URL="https://account-d.docusign.com/oauth/auth"

    params = {
        "client_id": settings.INTEGRATION_KEY,
        "redirect_uri": redirect_url,
        "scope": scope,  # Define the permissions you need
        "response_type": "code",
    }
    query = "&".join([f"{key}={value}" for key, value in params.items()])
    redirect_url = f"{AUTHORIZE_URL}?{query}"

    print("The redirect-url is: " + redirect_url)
    return redirect_url

async def fetch_agreements(agreement_id: str, nav_access_token):
    GET_ALL_NAV_AGREEMENTS_URL = f"https://api-d.docusign.com/v1/accounts/{settings.API_ACCOUNT_ID}/agreements"

    if agreement_id != "ALL":
        print(f"Id is present: {agreement_id}")
        GET_ALL_NAV_AGREEMENTS_URL = f"https://api-d.docusign.com/v1/accounts/{settings.API_ACCOUNT_ID}/agreements/{agreement_id}"
        print(f"The endpoint being called is: {GET_ALL_NAV_AGREEMENTS_URL}")

    async with httpx.AsyncClient() as client:
        try:
            response = await client.get(
                GET_ALL_NAV_AGREEMENTS_URL,
                headers={
                    "Authorization": f"Bearer {nav_access_token}",
                    "Accept": "application/json"
                },
            )
            response.raise_for_status()
        except httpx.HTTPError as e:
            print(f"HTTP error occurred: {e}")
            return {"error": str(e)}

    agreements = response.json()
    nav_agreements = agreements.get("data", [])

    if agreement_id != "ALL":
        nav_agreements = agreements

    if not nav_agreements:
        print("No agreements found or `data` key missing in response.")

    return nav_agreements

def get_envelope_status(envelope_id: str, sign_access_token: str):
    client = create_api_client(sign_access_token)
    envelopes_api = EnvelopesApi(client)
    envelope = envelopes_api.get_envelope(settings.API_ACCOUNT_ID, envelope_id)
    return envelope.status


def send_envelope(email: SignEmail, sign_access_token):
    try:
        # 1. Create the envelope request object
        envelope_definition = make_envelope(email)

        # 2. Call Envelopes::create API method
        api_client = create_api_client(sign_access_token)
        envelope_api = EnvelopesApi(api_client)
        results = envelope_api.create_envelope(settings.API_ACCOUNT_ID, envelope_definition=envelope_definition)
        return results.envelope_id

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to send envelope: {str(e)}")

def convert_text_to_pdf(text):
    buffer = BytesIO()
    c = canvas.Canvas(buffer)

    # Starting position for text
    x, y = 100, 750  # Adjust as needed for margins

    # Split text into lines based on \n
    lines = text.split("\n")
    for line in lines:
        c.drawString(x, y, line)
        y -= 15  # Move down for the next line (adjust spacing as needed)

        # Check if we're at the bottom of the page
        if y < 50:  # Prevent text from running off the page
            c.showPage()  # Start a new page
            y = 750  # Reset y position for the new page

    c.save()
    buffer.seek(0)
    return buffer.read()  # Binary PDF data

def make_envelope(email: SignEmail):
    # Check if file_content is plain-text, then convert to PDF
    if isinstance(email.file_content, str):
        pdf_content = convert_text_to_pdf(email.file_content)  # Convert plain-text to binary PDF
    else:
        pdf_content = email.file_content  # Assume it’s already binary PDF

    # Base64 encode the PDF content
    document = Document(
        document_base64=base64.b64encode(pdf_content).decode('utf-8'),
        name=email.file_name,
        file_extension="pdf",
        document_id="1"
    )

    # Define the signers
    signers = []
    for index, user in enumerate(email.primary_users):
        signer = Signer(
            email=user.email,
            name=f"{user.first_name} {user.last_name}",
            recipient_id=str(index + 1),
            routing_order=str(index + 1)
        )
        signers.append(signer)

    # Define CC users
    cc_offset = len(email.primary_users)
    # ccs = []
    # for index, cc_user in enumerate(email.cc_users):
    #     cc = CarbonCopy(
    #         email=cc_user.email,
    #         name=f"{cc_user.first_name} {cc_user.last_name}",
    #         recipient_id=str(index + cc_offset + 1),
    #         routing_order=str(index + cc_offset + 1)
    #     )
    #     ccs.append(cc)

    # Create envelope definition
    envelope_definition = EnvelopeDefinition(
        email_subject=email.subject,
        status="sent",
        documents=[document],
        recipients=Recipients(signers=signers)
        # recipients=Recipients(signers=signers, carbon_copies=ccs)
    )

    return envelope_definition

def create_api_client(sign_access_token):
    api_client = ApiClient()
    api_client.host = settings.DEV_BASE_PATH
    api_client.set_default_header(header_name="Authorization", header_value=f"Bearer {sign_access_token}")
    return api_client

def split_string_by_limit(input_string, char_limit):
    """
    Splits a string into a list of substrings, each with a length not exceeding the given character limit.

    Args:
        input_string (str): The input string to be split.
        char_limit (int): The maximum number of characters for each substring.

    Returns:
        list: A list of substrings.
    """
    if char_limit <= 0:
        raise ValueError("Character limit must be greater than 0.")

    words = input_string.split()
    result = []
    current_line = ""

    for word in words:
        # If adding the next word exceeds the limit, store the current line and start a new one
        if len(current_line) + len(word) + (1 if current_line else 0) > char_limit:
            result.append(current_line)
            current_line = word
        else:
            # Add the word to the current line
            current_line += (" " if current_line else "") + word

    # Add the last line if it exists
    if current_line:
        result.append(current_line)

    return result

def update_agreement(state: SectionState):
    """ Update a section of the report """

    # Get state
    agreement = state["agreement_text"]
    prompt = state["prompt"]
    provisions = state["provisions"]

    # Format system instructions
    system_instructions = magic_edit_instructions.format(prompt=prompt, agreement=agreement, provisions=provisions)

    # Update section
    section_content = llm.with_structured_output(AIEdit).invoke([SystemMessage(content=system_instructions)]+[HumanMessage(content="Update the given documents based on the provided input")])
    return {"response": section_content}


def find_differences(original_text, updated_text):
    lines1 = original_text.splitlines()
    lines2 = updated_text.splitlines()

    # Create a Differ object and compare the lines
    differ = Differ()
    diff = differ.compare(lines1, lines2)

    formatted_diffs = []
    for line in diff:
        if not line.startswith("?"):
            formatted_diffs.append(line)
        if line.strip() == "":
            # Replace spaces after the first one with '\n'
            formatted_diffs.append(line[0] + re.sub(r' +', '\n', line[1:]))

    return formatted_diffs

async def generate_report_plan(state: ReportState):
    """ Generate the report plan """
    # Inputs
    number_of_queries = state["number_of_queries"]
    tavily_topic = state["tavily_topic"]
    tavily_days = state.get("tavily_days", None)
    insight_type = state["insight_type"]
    agreement = state["agreement"]

    # Extract the given insight type
    extraction_instruction_query = insight_extraction_instructions.format(insight_type = insight_type, agreement = agreement)
    structured_extraction_llm = llm.with_structured_output(Extractions)
    extractions = structured_extraction_llm.invoke(extraction_instruction_query)

    # Generate web search queries
    structured_query_llm = llm.with_structured_output(Queries)
    system_instructions_query = report_planner_query_writer_instructions.format(insight_type = insight_type, insight_extraction = extractions, number_of_queries=number_of_queries)
    queries = structured_query_llm.invoke([SystemMessage(content=system_instructions_query)]+[HumanMessage(content="Generate search queries that will help with planning the sections of the report.")])
    query_list = [query.search_query for query in queries.queries]
    search_docs = await tavily_search_async(query_list, tavily_topic, tavily_days)
    source_str = deduplicate_and_format_sources(search_docs, max_tokens_per_source=1000, include_raw_content=False)

    # Generate the report structure
    system_instructions_sections = report_planner_instructions.format(insight_type = insight_type, insight_extraction = extractions, context=source_str)
    report_structured_llm = llm.with_structured_output(Insights)
    report_sections = report_structured_llm.invoke([SystemMessage(content=system_instructions_sections)]+[HumanMessage(content="Generate the sections of the report. Your response must include a 'insights' field containing a list of sections. Each section must have: type, explanation, extraction, document_lookup, deviation and insights_generated fields.")])

    # Update the state with insight outline
    return {"insights": report_sections.insights}

async def tavily_search_async(search_queries, tavily_topic, tavily_days):
    """
    Performs concurrent web searches using the Tavily API.

    Args:
        search_queries (List[SearchQuery]): List of search queries to process
        tavily_topic (str): Type of search to perform ('news' or 'general')
        tavily_days (int): Number of days to look back for news articles (only used when tavily_topic='news')

    Returns:
        List[dict]: List of search results from Tavily API, one per query

    Note:
        For news searches, each result will include articles from the last `tavily_days` days.
        For general searches, the time range is unrestricted.
    """
    tavily_async_client = AsyncTavilyClient(api_key= settings.TAVILY_API_KEY)
    search_tasks = []
    for query in search_queries:
        if tavily_topic == "news":
            search_tasks.append(
                tavily_async_client.search(
                    query,
                    max_results=5,
                    include_raw_content=True,
                    topic="news",
                    days=tavily_days
                )
            )
        else:
            search_tasks.append(
                tavily_async_client.search(
                    query,
                    max_results=5,
                    include_raw_content=True,
                    topic="general"
                )
            )

    # Execute all searches concurrently
    search_docs = await asyncio.gather(*search_tasks)

    return search_docs

def deduplicate_and_format_sources(search_response, max_tokens_per_source, include_raw_content=True):
    """
    Takes either a single search response or list of responses from Tavily API and formats them.
    Limits the raw_content to approximately max_tokens_per_source.
    include_raw_content specifies whether to include the raw_content from Tavily in the formatted string.

    Args:
        search_response: Either:
            - A dict with a 'results' key containing a list of search results
            - A list of dicts, each containing search results

    Returns:
        str: Formatted string with deduplicated sources
    """
    # Convert input to list of results
    if isinstance(search_response, dict):
        sources_list = search_response['results']
    elif isinstance(search_response, list):
        sources_list = []
        for response in search_response:
            if isinstance(response, dict) and 'results' in response:
                sources_list.extend(response['results'])
            else:
                sources_list.extend(response)
    else:
        raise ValueError("Input must be either a dict with 'results' or a list of search results")

    # Deduplicate by URL
    unique_sources = {}
    for source in sources_list:
        if source['url'] not in unique_sources:
            unique_sources[source['url']] = source

    # Format output
    formatted_text = "Sources:\n\n"
    for i, source in enumerate(unique_sources.values(), 1):
        formatted_text += f"Source {source['title']}:\n===\n"
        formatted_text += f"URL: {source['url']}\n===\n"
        formatted_text += f"Most relevant content from source: {source['content']}\n===\n"
        if include_raw_content:
            # Using rough estimate of 4 characters per token
            char_limit = max_tokens_per_source * 4
            # Handle None raw_content
            raw_content = source.get('raw_content', '')
            if raw_content is None:
                raw_content = ''
                print(f"Warning: No raw_content found for source {source['url']}")
            if len(raw_content) > char_limit:
                raw_content = raw_content[:char_limit] + "... [truncated]"
            formatted_text += f"Full source content limited to {max_tokens_per_source} tokens: {raw_content}\n\n"

    return formatted_text.strip()

def generate_insight(state: InsightState):
    """ Write a insights for a given clause or obligation"""

    # Get state
    insight = state["insight"]
    source_str = state["source_str"]


    # Format system instructions
    system_instructions = insight_writer_instructions.format(insight_type=insight.insight_type, extraction = insight.extraction, context=source_str)

    # Generate section
    section_content = llm.invoke([SystemMessage(content=system_instructions)]+[HumanMessage(content="Generate your insights based on the provided sources.")])

    # Write content to the section object
    insight_generated = section_content.content

    # Write the updated section to completed sections
    insight.insight_generated = insight_generated
    return {"completed_insights": [insight]}

async def search_web(state: InsightState):
    """ Search the web for each query, then return a list of raw sources and a formatted string of sources."""

    # Get state
    search_queries = state["search_queries"]

    # Get configuration
    tavily_topic = state["tavily_topic"]
    tavily_days = state["tavily_days"]

    # Web search
    query_list = [query.search_query for query in search_queries]
    search_docs = await tavily_search_async(query_list, tavily_topic, tavily_days)

    # Deduplicate and format sources
    source_str = deduplicate_and_format_sources(search_docs, max_tokens_per_source=5000, include_raw_content=True)

    return {"source_str": source_str}

def generate_queries(state: InsightState):
    """ Generate search queries for a insight section """

    # Get state
    insight = state["insight"]
    number_of_queries = state["number_of_queries"]

    # Generate queries
    structured_llm = llm.with_structured_output(Queries)

    # Format system instructions
    system_instructions = query_writer_instructions.format(insight_type=insight.insight_type, extraction = insight.extraction, number_of_queries=number_of_queries)

    # Generate queries
    queries = structured_llm.invoke([SystemMessage(content=system_instructions)]+[HumanMessage(content="Generate search queries on the provided topic.")])

    return {"search_queries": queries.queries}

def format_nav_extractions(nav_extractions: List):
    if not nav_extractions: return []
    formatted_documents = []

    for extraction in nav_extractions:
        document = {
            "document_id" : extraction["id"],
            "document_text": get_document_text(extraction["file_name"]),
            "navigator_extractions": extraction,
            "obligations": [],
            "signature_metadata": {"envelope_id": "", "signature_status": SignStatus.REVIEW.name},
            "clauses": [],
            "versions": [get_document_text(extraction["file_name"])]
        }

        formatted_documents.append(document)

    return formatted_documents


def get_document_text(file_name: str) -> str:
    """
    Recursively finds the document in the 'Sample Agreements' folder (including subfolders)
    and extracts its text.

    Args:
        file_name (str): The name of the file to extract text from.

    Returns:
        str: The extracted text from the file.

    Raises:
        FileNotFoundError: If the file is not found in the directory.
        ValueError: If the file type is unsupported.
    """
    # Define the root folder path
    folder_path = Path("./SampleAgreements")

    # Check if the folder exists
    if not folder_path.exists():
        raise FileNotFoundError(f"The directory '{folder_path}' does not exist.")

    # Recursively search for the file
    file_path = None
    for path in folder_path.rglob(file_name):
        if path.is_file() and path.name == file_name:
            file_path = path
            break

    file_found = True
    # If file is not found
    if file_path is None:
        file_found = False
        print("File not found")
        return("Content Unknown")
        # raise FileNotFoundError(f"The file '{file_name}' was not found in '{folder_path}' or its subdirectories.")

    if file_found:
        # Extract text based on file type
        if file_path.suffix == ".txt":
            with open(file_path, "r", encoding="utf-8") as file:
                return file.read()
        elif file_path.suffix == ".pdf":
            import PyPDF2
            with open(file_path, "rb") as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text()
                return text
        elif file_path.suffix == ".docx":
            from docx import Document
            doc = Document(file_path)
            return "\n".join(paragraph.text for paragraph in doc.paragraphs)
        else:
            raise ValueError(f"Unsupported file type: {file_path.suffix}")